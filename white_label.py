#!/usr/bin/env python3
"""
Astoria White-Label Tool
========================
Batch-replace the Astoria logo in monthly commentary DOCX files
with client logos and export as PDF.

Usage:
    python white_label.py --input commentary.docx --logos ./client_logos/ --output ./output/
    python white_label.py --input commentary.docx --logos ./client_logos/ --output ./output/ --format docx
    python white_label.py --input commentary.docx --logos ./client_logos/ --output ./output/ --format both

Requirements:
    pip install python-docx Pillow
    LibreOffice (for PDF conversion): sudo apt install libreoffice
"""

import argparse
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image


# Supported image formats for client logos
SUPPORTED_LOGO_FORMATS = {'.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff', '.webp'}


def find_logo_image(doc: Document) -> tuple[str, str]:
    """
    Find the Astoria logo in the document.
    
    Strategy: Look for the first image in the first paragraph of the document body.
    The Astoria commentary always places the logo in paragraph 0 as an inline drawing.
    
    Returns:
        tuple of (relationship_id, image_target_path) e.g. ('rId9', 'media/image1.jpeg')
    """
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    }
    
    # Search first 3 paragraphs for an image
    for para_idx in range(min(3, len(doc.paragraphs))):
        para = doc.paragraphs[para_idx]
        drawings = para._element.findall('.//w:drawing', nsmap)
        
        for drawing in drawings:
            blips = drawing.findall('.//a:blip', nsmap)
            for blip in blips:
                embed_id = blip.get(
                    '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                )
                if embed_id and embed_id in doc.part.rels:
                    rel = doc.part.rels[embed_id]
                    target = rel.target_ref
                    # Check if it's an image (not a chart or other embedded object)
                    if any(ext in target.lower() for ext in ['.jpeg', '.jpg', '.png', '.bmp', '.gif', '.emf']):
                        return embed_id, target
    
    raise ValueError(
        "Could not find the Astoria logo in the document. "
        "Make sure the DOCX has an image in the first few paragraphs."
    )


def replace_logo_in_docx(input_path: str, logo_path: str, output_path: str) -> None:
    """
    Replace the Astoria logo in a DOCX with a client logo.
    
    This works by:
    1. Opening the DOCX (which is a ZIP file)
    2. Finding the logo image relationship
    3. Replacing the image binary data with the new client logo
    4. Saving as a new DOCX
    """
    import zipfile
    import io
    
    # First, identify which image file is the logo
    doc = Document(input_path)
    rel_id, image_target = find_logo_image(doc)
    
    # The image_target is like 'media/image1.jpeg' - we need 'word/media/image1.jpeg'
    image_zip_path = f"word/{image_target}"
    
    # Determine the format the original image is in
    original_ext = Path(image_target).suffix.lower()
    
    # Prepare the replacement logo
    # We need to convert it to match the original format to avoid compatibility issues
    logo_img = Image.open(logo_path)
    
    # Convert to RGB if necessary (for JPEG)
    if original_ext in ('.jpg', '.jpeg'):
        if logo_img.mode in ('RGBA', 'P'):
            # Create white background for transparent logos
            background = Image.new('RGB', logo_img.size, (255, 255, 255))
            if logo_img.mode == 'P':
                logo_img = logo_img.convert('RGBA')
            background.paste(logo_img, mask=logo_img.split()[3] if logo_img.mode == 'RGBA' else None)
            logo_img = background
        elif logo_img.mode != 'RGB':
            logo_img = logo_img.convert('RGB')
        
        img_buffer = io.BytesIO()
        logo_img.save(img_buffer, format='JPEG', quality=95)
    elif original_ext == '.png':
        img_buffer = io.BytesIO()
        logo_img.save(img_buffer, format='PNG')
    else:
        # For other formats, save as PNG
        img_buffer = io.BytesIO()
        logo_img.save(img_buffer, format='PNG')
    
    replacement_bytes = img_buffer.getvalue()
    
    # Now rebuild the DOCX ZIP with the replaced image
    with zipfile.ZipFile(input_path, 'r') as zin:
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == image_zip_path:
                    # Replace with our client logo
                    zout.writestr(item, replacement_bytes)
                else:
                    zout.writestr(item, data)


def convert_docx_to_pdf(docx_path: str, output_dir: str) -> str:
    """
    Convert a DOCX file to PDF using LibreOffice.
    
    Returns:
        Path to the generated PDF file.
    """
    result = subprocess.run(
        [
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', output_dir,
            docx_path
        ],
        capture_output=True,
        text=True,
        timeout=120
    )
    
    if result.returncode != 0:
        raise RuntimeError(
            f"LibreOffice conversion failed:\n{result.stderr}"
        )
    
    # The output PDF will have the same name as the DOCX but with .pdf extension
    pdf_name = Path(docx_path).stem + '.pdf'
    pdf_path = os.path.join(output_dir, pdf_name)
    
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(
            f"Expected PDF at {pdf_path} but it was not created. "
            f"LibreOffice output: {result.stdout}"
        )
    
    return pdf_path


def get_client_name_from_logo(logo_path: str) -> str:
    """
    Extract a clean client name from the logo filename.
    
    Examples:
        'Bluedoor_Logo.png' -> 'Bluedoor'
        'acme-financial.jpg' -> 'Acme Financial'
        'logo_smith_wealth.png' -> 'Smith Wealth'
    """
    stem = Path(logo_path).stem
    # Remove common prefixes/suffixes
    for remove in ['logo', 'Logo', 'LOGO', '_logo', '-logo', 'logo_', 'logo-']:
        stem = stem.replace(remove, '')
    # Clean up separators
    stem = stem.replace('_', ' ').replace('-', ' ').strip()
    # Title case
    return stem.title() if stem else Path(logo_path).stem


def process_batch(
    input_docx: str,
    logos_dir: str,
    output_dir: str,
    output_format: str = 'pdf',
    naming_template: str = '{month}_{year}_Monthly_Commentary_Report_{client}'
) -> list[dict]:
    """
    Process a batch of client logos against a single commentary DOCX.
    
    Args:
        input_docx: Path to the Astoria commentary DOCX
        logos_dir: Directory containing client logo image files
        output_dir: Directory to save branded outputs
        output_format: 'pdf', 'docx', or 'both'
        naming_template: Template for output filenames
    
    Returns:
        List of dicts with processing results per client
    """
    # Validate inputs
    if not os.path.exists(input_docx):
        raise FileNotFoundError(f"Input DOCX not found: {input_docx}")
    if not os.path.isdir(logos_dir):
        raise NotADirectoryError(f"Logos directory not found: {logos_dir}")
    
    os.makedirs(output_dir, exist_ok=True)
    
    # Collect all logo files
    logo_files = sorted([
        os.path.join(logos_dir, f)
        for f in os.listdir(logos_dir)
        if Path(f).suffix.lower() in SUPPORTED_LOGO_FORMATS
    ])
    
    if not logo_files:
        raise ValueError(
            f"No logo files found in {logos_dir}. "
            f"Supported formats: {', '.join(SUPPORTED_LOGO_FORMATS)}"
        )
    
    # Extract month/year from the input filename for naming
    input_stem = Path(input_docx).stem
    # Try to extract month and year from filename
    import re
    month_match = re.search(
        r'(January|February|March|April|May|June|July|August|September|October|November|December)',
        input_stem, re.IGNORECASE
    )
    year_match = re.search(r'(20\d{2})', input_stem)
    month = month_match.group(1) if month_match else 'Monthly'
    year = year_match.group(1) if year_match else ''
    
    results = []
    total = len(logo_files)
    
    # Create a temp directory for intermediate DOCX files
    with tempfile.TemporaryDirectory() as tmp_dir:
        for i, logo_path in enumerate(logo_files, 1):
            client_name = get_client_name_from_logo(logo_path)
            
            # Build output filename
            base_name = naming_template.format(
                month=month, year=year, client=client_name.replace(' ', '_')
            )
            
            result = {
                'client': client_name,
                'logo': logo_path,
                'status': 'pending',
                'outputs': []
            }
            
            try:
                print(f"  [{i}/{total}] Processing: {client_name}...", end=' ', flush=True)
                
                # Step 1: Replace logo in DOCX
                if output_format in ('docx', 'both'):
                    docx_output = os.path.join(output_dir, f"{base_name}.docx")
                else:
                    docx_output = os.path.join(tmp_dir, f"{base_name}.docx")
                
                replace_logo_in_docx(input_docx, logo_path, docx_output)
                
                if output_format in ('docx', 'both'):
                    result['outputs'].append(docx_output)
                
                # Step 2: Convert to PDF if needed
                if output_format in ('pdf', 'both'):
                    pdf_path = convert_docx_to_pdf(docx_output, output_dir)
                    result['outputs'].append(pdf_path)
                    
                    # Clean up intermediate DOCX if only PDF was requested
                    if output_format == 'pdf' and os.path.exists(docx_output) and tmp_dir in docx_output:
                        pass  # tempdir will clean it up
                
                result['status'] = 'success'
                print("✓")
                
            except Exception as e:
                result['status'] = 'error'
                result['error'] = str(e)
                print(f"✗ ({e})")
            
            results.append(result)
    
    return results


def main():
    parser = argparse.ArgumentParser(
        description='Astoria White-Label Tool — Batch brand monthly commentaries with client logos',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s --input commentary.docx --logos ./client_logos/ --output ./branded/
  %(prog)s --input commentary.docx --logos ./client_logos/ --output ./branded/ --format both
  %(prog)s --input commentary.docx --logos ./client_logos/ --output ./branded/ --format docx
        """
    )
    
    parser.add_argument(
        '--input', '-i',
        required=True,
        help='Path to the Astoria monthly commentary DOCX file'
    )
    parser.add_argument(
        '--logos', '-l',
        required=True,
        help='Directory containing client logo image files (PNG, JPG, etc.)'
    )
    parser.add_argument(
        '--output', '-o',
        default='./output',
        help='Output directory for branded files (default: ./output)'
    )
    parser.add_argument(
        '--format', '-f',
        choices=['pdf', 'docx', 'both'],
        default='pdf',
        help='Output format (default: pdf)'
    )
    parser.add_argument(
        '--naming', '-n',
        default='{month}_{year}_Monthly_Commentary_Report_{client}',
        help='Naming template for output files (default: {month}_{year}_Monthly_Commentary_Report_{client})'
    )
    
    args = parser.parse_args()
    
    print("=" * 60)
    print("  Astoria White-Label Tool")
    print("=" * 60)
    print(f"  Input:   {args.input}")
    print(f"  Logos:   {args.logos}")
    print(f"  Output:  {args.output}")
    print(f"  Format:  {args.format}")
    print("=" * 60)
    
    # Count logos
    logo_count = len([
        f for f in os.listdir(args.logos)
        if Path(f).suffix.lower() in SUPPORTED_LOGO_FORMATS
    ])
    print(f"\n  Found {logo_count} client logos. Starting batch processing...\n")
    
    results = process_batch(
        input_docx=args.input,
        logos_dir=args.logos,
        output_dir=args.output,
        output_format=args.format,
        naming_template=args.naming
    )
    
    # Summary
    success = sum(1 for r in results if r['status'] == 'success')
    errors = sum(1 for r in results if r['status'] == 'error')
    
    print(f"\n{'=' * 60}")
    print(f"  COMPLETE: {success} succeeded, {errors} failed")
    print(f"  Output directory: {os.path.abspath(args.output)}")
    print(f"{'=' * 60}")
    
    if errors > 0:
        print("\n  Errors:")
        for r in results:
            if r['status'] == 'error':
                print(f"    - {r['client']}: {r['error']}")
    
    return 0 if errors == 0 else 1


if __name__ == '__main__':
    sys.exit(main())
